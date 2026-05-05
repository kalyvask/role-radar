"""Render a PrepDoc to a .docx file using python-docx.

Stylistically mirrors the Snowflake reference doc — clean sectioned layout,
headings, body text with inline bold/italic from the Markdown.

Markdown handling is intentionally lightweight: paragraphs separated by blank
lines, `-`/`*` bullets, `**bold**` and `*italic*` inline. Code blocks (triple
backticks) become monospace paragraphs. Anything fancier degrades gracefully to
plain text.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from role_radar.interview_prep.models import (
    BackgroundFit,
    Closing,
    InterviewQuestion,
    PrepDay,
    PrepDoc,
    TopicToKnow,
)


class DocxRenderError(RuntimeError):
    """Raised when the DOCX renderer fails."""


def render_docx(doc: PrepDoc, output_path: Path) -> Path:
    """Render a PrepDoc to a .docx file. Returns the output path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise DocxRenderError(
            "python-docx is required for DOCX rendering. Install with `pip install python-docx`."
        ) from e

    document = Document()

    # Default style
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h = doc.header

    # Title
    title = document.add_heading(f"{h.company} — {h.role_title}", level=0)
    title.alignment = 1  # center

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    subtitle_run = subtitle.add_run("Interview Preparation Document")
    subtitle_run.bold = True

    meta_bits = [h.location]
    if h.compensation:
        meta_bits.append(h.compensation)
    meta = document.add_paragraph(" | ".join(meta_bits))
    meta.alignment = 1

    gen_date = doc.generated_at or datetime.utcnow()
    candidate_line = document.add_paragraph()
    candidate_line.alignment = 1
    candidate_run = candidate_line.add_run(f"{h.candidate_name} · {gen_date.strftime('%B %Y')}")
    candidate_run.italic = True
    candidate_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_paragraph()  # spacer

    # Section 1
    document.add_heading("1  /  What this role actually is", level=1)
    _render_markdown_block(document, doc.role_summary_markdown)

    # Section 2
    document.add_heading("2  /  Topics you must know cold", level=1)
    document.add_paragraph(
        "These are the topics where being unprepared will end the interview in "
        "the first five minutes. Each is summarized to the level you need to "
        "discuss intelligently with the hiring manager and engineering partners."
    )
    for topic in doc.topics_to_know:
        _render_topic(document, topic)

    # Section 3 (optional)
    section_n = 3
    if doc.strategic_angle_markdown:
        document.add_heading("3  /  The strategic angle", level=1)
        _render_markdown_block(document, doc.strategic_angle_markdown)
        section_n = 4

    # Section 4
    document.add_heading(f"{section_n}  /  Likely interview questions", level=1)
    for q in doc.likely_questions:
        _render_question(document, q)

    # Section 5
    section_n += 1
    document.add_heading(f"{section_n}  /  How your background maps — and the gaps", level=1)
    _render_background_fit(document, doc.background_fit)

    # Section 6
    section_n += 1
    document.add_heading(f"{section_n}  /  The preparation plan", level=1)
    for day in doc.prep_plan:
        _render_prep_day(document, day)

    # Section 7
    section_n += 1
    document.add_heading(f"{section_n}  /  Closing", level=1)
    _render_closing(document, doc.closing)

    # Sources
    if doc.sources_markdown:
        document.add_heading("Sources", level=1)
        _render_markdown_block(document, doc.sources_markdown)

    document.save(str(output_path))
    return output_path


def _render_topic(document, topic: TopicToKnow) -> None:
    document.add_heading(f"{topic.section_number}  {topic.title}", level=2)
    _render_markdown_block(document, topic.body_markdown)


def _render_question(document, q: InterviewQuestion) -> None:
    document.add_heading(f"{q.section_number}  {q.category}", level=2)
    p = document.add_paragraph()
    p.add_run("Q: ").bold = True
    p.add_run(q.question)

    p = document.add_paragraph()
    p.add_run("Approach: ").bold = True
    p.add_run(q.approach)

    _render_markdown_block(document, q.sample_answer_markdown)

    why = document.add_paragraph()
    why_run = why.add_run(f"Why this works: {q.why_this_works}")
    why_run.italic = True


def _render_background_fit(document, bf: BackgroundFit) -> None:
    document.add_heading("What you bring that maps directly", level=2)
    for item in bf.what_you_bring:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("The honest gaps", level=2)
    for item in bf.honest_gaps:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("How to frame the gap honestly", level=2)
    _render_markdown_block(document, bf.how_to_frame_gap_markdown)


def _render_prep_day(document, day: PrepDay) -> None:
    document.add_heading(f"{day.day} — {day.focus}", level=2)
    _render_markdown_block(document, day.items_markdown)


def _render_closing(document, c: Closing) -> None:
    document.add_heading("Three questions to ask, ranked by signal", level=2)
    for i, q in enumerate(c.questions_to_ask, 1):
        document.add_paragraph(f"{i}. {q}")

    document.add_heading("Disqualifiers for this role", level=2)
    for d in c.disqualifiers:
        document.add_paragraph(d, style="List Bullet")

    document.add_heading("The single strongest move you can make", level=2)
    _render_markdown_block(document, c.strongest_single_move_markdown)

    if c.final_note_markdown:
        document.add_heading("Final note", level=2)
        _render_markdown_block(document, c.final_note_markdown)


# ──────────────────────────────────────────────────────────────────────
# Lightweight Markdown → DOCX paragraph rendering.
# Handles paragraphs, bullet lists, and inline **bold** / *italic*.
# ──────────────────────────────────────────────────────────────────────

_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_HEADING_RE = re.compile(r"^\s*(#{1,6})\s+(.*)$")
_CODE_FENCE_RE = re.compile(r"^\s*```")
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_|`[^`]+`)")


def _render_markdown_block(document, text: str) -> None:
    """Render a block of Markdown into the document.

    Supports paragraphs, bullet/numbered lists, fenced code blocks (rendered
    as monospace), and inline `**bold**` / `*italic*` / `_italic_` / `` `code` ``.
    """
    if not text:
        return

    lines = text.strip().splitlines()
    in_code = False
    code_buffer: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        joined = " ".join(line.strip() for line in paragraph_buffer if line.strip())
        if joined:
            p = document.add_paragraph()
            _add_runs_with_inline_formatting(p, joined)
        paragraph_buffer.clear()

    def flush_code() -> None:
        if not code_buffer:
            return
        p = document.add_paragraph()
        run = p.add_run("\n".join(code_buffer))
        run.font.name = "Consolas"
        code_buffer.clear()

    for line in lines:
        if _CODE_FENCE_RE.match(line):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        bullet_match = _BULLET_RE.match(line)
        numbered_match = _NUMBERED_RE.match(line)
        heading_match = _HEADING_RE.match(line)

        if bullet_match:
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            _add_runs_with_inline_formatting(p, bullet_match.group(1))
        elif numbered_match:
            flush_paragraph()
            p = document.add_paragraph(style="List Number")
            _add_runs_with_inline_formatting(p, numbered_match.group(1))
        elif heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)) + 1, 4)  # nest under section heading
            document.add_heading(heading_match.group(2), level=level)
        elif not line.strip():
            flush_paragraph()
        else:
            paragraph_buffer.append(line)

    flush_paragraph()
    flush_code()


def _add_runs_with_inline_formatting(paragraph, text: str) -> None:
    """Split text on inline Markdown markers and add runs with the right formatting."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)
